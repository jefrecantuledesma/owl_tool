#!/bin/pypy3
from docx import Document
#from docx.shared import Pt # Broken feature, to be fixed later.
from docx.enum.text import WD_BREAK
import sys
import os.path
from datetime import date

config_path = os.path.expanduser("~/.config/owl_tool/config")
does_config_exist = os.path.exists(config_path)

def request_download_path() -> str:
    path: str = input("Please input the path of the downloaded document: ").replace("'", "").strip()
    if path[0] == "~":
        path = os.path.expanduser(path)
    if os.path.exists(path) == False:
        print("That path does not exist.")
        return request_download_path()
    return path

def get_download_path() -> str:
    path: str = ""
    for i in range(0, len(sys.argv)):
        if sys.argv[i].replace("-", "").strip() == "d" or sys.argv[i].replace("-", "").strip() == "download":
            path: str = sys.argv[i + 1]
            if path[0] == "~":
                path = os.path.expanduser(path)
            if os.path.exists(path) == False:
                print("That path does not exist.")
                return request_download_path()
            return path
    return request_download_path()

def request_template_path() -> str:
    path: str = input("Please input the path of the template document: ").replace("'", "").strip()
    if path[0] == "~":
        path = os.path.expanduser(path)
    if os.path.exists(path) == False:
        print("That path does not exist.")
        return request_template_path()
    return path

def get_template_path() -> str:
    path: str = ""
    for i in range(0, len(sys.argv)):
        if sys.argv[i].replace("-", "").strip() == "t" or sys.argv[i].replace("-", "").strip() == "template":
            path: str = sys.argv[i + 1]
            if path[0] == "~":
                path = os.path.expanduser(path)
            if os.path.exists(path) == False:
                print("That path does not exist.")
                return request_template_path()
            return path
    return request_template_path()

def select_editable_region(template_document) -> str:
    all_paras = template_document.paragraphs
    start_para: int = 0 
    end_para: int = 0
    counter: int = 0
    for para in all_paras:
        if para.text.startswith("#"):
            start_para=counter
        if para.text.endswith("#"):
            end_para=counter
        counter += 1
    template_list: list[str] = []
    for index in range(start_para, end_para + 1):
        template_list.append(all_paras[index].text.replace("#", ""))
    template_text = "\n".join(template_list)
    return template_text

def edit_regions(edit_text, begin, end) -> str:
    edit_text = edit_text.replace("<", "").replace(">", "")
    new_text = input(f"Here's the editable text: {edit_text}.\nEnter what would you like to replace it with: ").strip()
    return new_text

def find_edit_regions(template_text) -> str:
    filled_template = template_text
    for begin in range(0, len(template_text)):
        if template_text[begin - 1] == "<":
            for end in range(begin, len(template_text)):
                if template_text[end] == ">":
                    edit_text = template_text[begin:end+1]
                    new_text = edit_regions(edit_text, begin, end)
                    filled_template = filled_template.replace(edit_text, new_text, 1)
                    break
    filled_template = filled_template.replace("<", "")
    filled_template = filled_template.replace(">", "")
    return filled_template

def request_save_path() -> str:
    path: str = input("Please input the path to save your document (leave blank for default): ").replace("'", "").strip()
    if path == "":
        path =  "~/Documents/"
    if path[0] == "~":
        path = os.path.expanduser(path)
    if path [-1] != "/":
        path = path + "/"
    if os.path.exists(path) == False:
        print("That path does not exist.")
        return request_save_path()
    return path

def get_save_path() -> str:
    path: str = ""
    for i in range(0, len(sys.argv)):
        if sys.argv[i].replace("-", "").strip() == "s" or sys.argv[i].replace("-", "").strip() == "save":
            path: str = sys.argv[i + 1]
            if path[0] == "~":
                path = os.path.expanduser(path)
            if os.path.exists(path) == False:
                print("That path does not exist.")
                return request_save_path()
            return path
    return request_save_path()

def write_text(edited_text, downloaded_document) -> None:
    first_paragraph = downloaded_document.paragraphs[0]
    #style = downloaded_document.styles['Normal'] # Unworking feature, fix later.
    #style.font.name = 'Times New Roman'
    #style.font.size = Pt(12)
    _ = first_paragraph.insert_paragraph_before(edited_text, style = downloaded_document.styles['Normal'])
    downloaded_document.paragraphs[0] = downloaded_document.styles["Normal"]
    for paragraph in downloaded_document.paragraphs:
        if edited_text in paragraph.text:
            run = paragraph.add_run()
            run.add_break(WD_BREAK.PAGE)
    today = str(date.today()).replace("-", ".")
    first_name = input("Please enter your first name: ").strip().lower()
    writer_lname = input("Please enter the writer's last name: ").strip().lower()
    save_path = get_save_path()
    downloaded_document.save(f"{save_path}{first_name}_OWL_{writer_lname}_{today}.docx")

def parse_config():
    settings: list = []
    with open(config_path, "r") as config:
        for line in config:
            if ("download" in line):
                settings.append(line.replace("download = ", "").strip())
                print(settings)


def main():
    if does_config_exist:
        parse_config()
    else:
        download_path: str = get_download_path()
        downloaded_document = Document(download_path)
        template_path: str = get_template_path()
        template_document = Document(template_path)
        template_text = select_editable_region(template_document)
        edited_text = find_edit_regions(template_text)
        write_text(edited_text, downloaded_document)

if __name__ == "__main__":
    main()
